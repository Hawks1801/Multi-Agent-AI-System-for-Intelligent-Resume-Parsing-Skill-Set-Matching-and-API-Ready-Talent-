import json
import random
import os
from docx import Document

def generate_resumes():
    print("--- Generating 500+ Synthetic Resumes (DOCX & TXT) ---")
    
    # Load base skills for relevance
    with open("data/taxonomy/skills_large.json", "r") as f:
        taxonomy = json.load(f)
    
    all_tech_skills = []
    def get_skills(data):
        if isinstance(data, list): all_tech_skills.extend(data)
        elif isinstance(data, dict):
            for v in data.values(): get_skills(v)
    get_skills(taxonomy["Technology"])

    names = ["Alice", "Bob", "Charlie", "David", "Eve", "Frank", "Grace", "Heidi", "Ivan", "Judy"]
    last_names = ["Smith", "Johnson", "Williams", "Brown", "Jones", "Garcia", "Miller", "Davis", "Rodriguez", "Martinez"]
    locations = ["New York", "San Francisco", "Austin", "Seattle", "Chicago", "Remote", "London", "Berlin", "Toronto", "Sydney"]

    os.makedirs("data/samples/resumes", exist_ok=True)
    
    ground_truth = []

    for i in range(520):
        first = random.choice(names)
        last = random.choice(last_names)
        full_name = f"{first} {last}"
        email = f"{first.lower()}.{last.lower()}{i}@example.com"
        location = random.choice(locations)
        
        # Pick random skills
        num_skills = random.randint(5, 12)
        resume_skills = random.sample(all_tech_skills, num_skills)
        
        # Record ground truth
        resume_data = {
            "id": f"resume_{i+1}",
            "name": full_name,
            "skills": resume_skills
        }
        ground_truth.append(resume_data)

        content = f"""{full_name}
Email: {email}
Location: {location}

Summary
Experienced professional with expertise in {', '.join(resume_skills[:3])}.

Skills
{', '.join(resume_skills)}

Experience
Lead Developer | Tech Corp | 2018-Present
- Managed a team of engineers using {resume_skills[0]}.
- Implemented solutions in {resume_skills[1]}.

Education
BS Computer Science | Tech University | 2014-2018
"""
        
        # Alternate formats
        if i % 2 == 0:
            # TXT
            with open(f"data/samples/resumes/resume_{i+1}.txt", "w") as f:
                f.write(content)
        else:
            # DOCX
            doc = Document()
            doc.add_heading(full_name, 0)
            doc.add_paragraph(f"Email: {email} | Location: {location}")
            doc.add_heading("Summary", level=1)
            doc.add_paragraph(f"Experienced professional with expertise in {', '.join(resume_skills[:3])}.")
            doc.add_heading("Skills", level=1)
            doc.add_paragraph(', '.join(resume_skills))
            doc.add_heading("Experience", level=1)
            doc.add_paragraph(f"Lead Developer | Tech Corp | 2018-Present")
            doc.add_paragraph(f"Managed a team of engineers using {resume_skills[0]}.")
            doc.save(f"data/samples/resumes/resume_{i+1}.docx")
            
        if (i+1) % 100 == 0:
            print(f"Generated {i+1} resumes...")

    with open("data/samples/ground_truth_resumes.json", "w") as f:
        json.dump(ground_truth, f, indent=2)

    print(f"Successfully generated {len(ground_truth)} resumes in data/samples/resumes/")

if __name__ == "__main__":
    generate_resumes()
